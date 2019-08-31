import docx
from docx import Document
import xlrd
import os
def read():
    loc = ("C:\\01\\Image processing\\Excel Test\\Book1.xlsx") 
    rd=[]
    wb = xlrd.open_workbook(loc) 
    sheet = wb.sheet_by_index(0) 
    sheet.cell_value(0, 0) 
    for i in range(sheet.nrows): 
    	rd.append(sheet.cell_value(i, 0))
    return rd	
def map(l):
    print(1,l[0],end=" | ")
    for i in range(1,len(l)):
        if(i%6!=0 and i!=0):
             print(i+1,l[i],end=" | ")
        else:
            print(i+1,l[i])        
def remove(l):
       print("How many vehicles do you have to remove?")
       No_Veh=int(input())
       i=0
       while(i!=No_Veh):
            print("The position of vehicle",i+1,"to remove?")
            r=int(input())
            try: 
               print("The chasis no at the given position is",l[r-1])
               print("Remove? press Y or N")
               if(input()=='Y'):
                   l[r-1]=0
                   IO[1]+=1
               i+=1    
            except:
                print("Enter a valid location")
       text2=str()    
       document2=Document()
       for i in range(0,len(l)):
          if(i!=len(l)-1):
            text2=text2+str(l[i])+" "
          else:
            text2=text2+str(l[i])
       para1=document2.add_paragraph(text2)
       #print(text2)      
       document2.save('C:\\01\\Image processing\\Backup\\bup.docx')
       text3=str()    
       document3=Document()
       for i in range(0,len(IO)):
          if(i!=len(IO)-1):
            text3=text3+str(IO[i])+" "
          else:
            text3=text3+str(IO[i])
       para1=document3.add_paragraph(text3)
       #print(text2)      
       document3.save('C:\\01\\Image processing\\Backup\\bup2.docx')
       return l,IO
       return l
def search(l):
    print("Enter the number of vehicles you wanna search?")
    No_Veh=int(input())
    op=[]
    for i in range(0,No_Veh):
        print("Enter Chasis No",i+1) 
        op.append(input())
    for i in range(0,No_Veh):
        if op[i] in l:
            x=l.index(op[i])
            print(op[i],"is present in",x+1)
        else:
            print(op[i],"is not present")
def reallocate(l):
    print("Enter the location from which you wish to move the vehicle")
    l1=int(input())
    print("Enter the location to which you want to move the vehicle")
    l2=int(input())
    try:
       temp=l[l2-1]
       l[l2-1]=l[l1-1]
       l[l1-1]=temp
       #print(l)
       print("Reallocation complete")
    except:
        print("Please Enter valid Location")
    document2=Document()
    text2=str()
    for i in range(0,len(l)):
          if(i!=len(l)-1):
            text2=text2+str(l[i])+" "
          else:
            text2=text2+str(l[i])
    para1=document2.add_paragraph(text2)
        #print(text2)      
    document2.save('C:\\01\\Image processing\\Backup\\bup.docx')
    return l
#x=os.listdir('/01/Image processing/Backup/')
document= Document('/01/Image processing/Backup/bup.docx')
for para in document.paragraphs:
    temp=para.text
    #print(temp)
l=list(temp.split(" "))
document3= Document('/01/Image processing/Backup/bup2.docx')
for para in document3.paragraphs:
    temp3=para.text
    #print(temp3)
IO=list(temp3.split(" "))
IO[0]=int(IO[0])
IO[1]=int(IO[1])
#print(l)    
while(1):
   print("What do you want to do?")
   print("1. Inward New vehicle")
   print("2. Map")  
   print("3. Remove")
   print("4. Search")
   print("5. Reallocate")
   do=int(input())
   if(do==1):
      rd=[]
      rd=read()
      flag=0
      for i in range(0,len(rd)):
       if rd[i] not in l:   
        print("Is it a")
        print("1. Longer Wheel Base Vehicle")
        print("2. Shorter Wheel Base Vehicle")
        Type=int(input())
        if(Type==1):
            print("Enter position between 1 to 25")
            flag=1
        elif(Type==2):
            print("Enter position between 26 to 50")
            flag=1
        if(flag==1):    
            get=int(input())
            l[get-1]=rd[i]
            IO[0]+=1
            text2=str()
            document2=Document()
            for i in range(0,len(l)):
              if(i!=len(l)-1):
                text2=text2+str(l[i])+" "
              else:
                text2=text2+str(l[i])
            para1=document2.add_paragraph(text2)
            #print(text2)      
            document2.save('C:\\01\\Image processing\\Backup\\bup.docx')
            text3=str()
            document4=Document()
            for i in range(0,len(IO)):
              if(i!=len(IO)-1):
                text3=text3+str(IO[i])+" "
              else:
                text3=text3+str(IO[i])
            para1=document4.add_paragraph(text3)
            #print(text2)      
            document4.save('C:\\01\\Image processing\\Backup\\bup2.docx')
            print("Vehicle Inwarded")
        else:
            print("Enter a valid number")
   if(do==2):
       map(l)
       print("")    
   if(do==3):
       l,IO=remove(l)
       print("Vehicle removed successully")
   if(do==4):
       search(l)
   if(do==5):
       l=reallocate(l)
       
