import docx
from docx import Document
import xlrd
import os
import xlwt 
from xlwt import Workbook 
x=os.listdir('/01/Image processing/Backup/')
document= Document('/01/Image processing/Backup/bup.docx')
for para in document.paragraphs:
    temp=para.text
l=list(temp.split(" "))
document2= Document('/01/Image processing/Backup/bup2.docx')
for para in document2.paragraphs:
    temp2=para.text
IO=list(temp2.split(" "))
wb=Workbook()
sheet1=wb.add_sheet('Sheet 1')
sheet1.write(0,0,'Location')
sheet1.write(0,1,'Chasis Number')
for i in range(0,len(l)):
    sheet1.write(i+1,0,i+1)
    sheet1.write(i+1,1,l[i])
sheet1.write(len(l)+1,0,"No of Inwarded Vehicles")
sheet1.write(len(l)+1,1,IO[0])
sheet1.write(len(l)+2,0,"No of Outwarded Vehicles")
sheet1.write(len(l)+2,1,IO[1])
wb.save('Report.xls')
IO[0]=0
IO[1]=0
document3= Document()
text=str(IO[0])+" "+str(IO[1])
para1=document3.add_paragraph(text)
document3.save('C:\\01\\Image processing\\Backup\\bup2.docx')

